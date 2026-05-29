"""简历导出：生成 Word (.docx) 和 PDF 文件。"""

from __future__ import annotations

import io
import os
import platform
from collections import defaultdict
from pathlib import Path

from career_store import CandidateProfile, ProfileEvidence

# ---- 常量 ----

FONT_CN = "Microsoft YaHei"
FONT_EN = "Calibri"
TEXT_COLOR = (28, 41, 56)
MUTED_COLOR = (86, 101, 115)
ACCENT_COLOR = (25, 89, 125)

CATEGORY_LABELS = {
    "education": "教育背景",
    "skill": "专业技能",
    "project": "项目经历",
    "award": "获奖荣誉",
    "availability": "求职意向",
}

CATEGORY_ORDER = ["education", "skill", "project", "award", "availability"]


def _group_evidence(evidence: list[ProfileEvidence]) -> dict[str, list[ProfileEvidence]]:
    groups: dict[str, list[ProfileEvidence]] = defaultdict(list)
    for e in evidence:
        groups[_field(e, "category")].append(e)
    return groups


def _field(item: object, name: str, default: str = "") -> str:
    if isinstance(item, dict):
        value = item.get(name, default)
    else:
        value = getattr(item, name, default)
    return default if value is None else str(value)


# ============================================================
# Word (.docx) 导出
# ============================================================

def _docx_set_run(run, size=None, color=TEXT_COLOR, bold=None, italic=None):
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = FONT_CN
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_CN)
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _docx_bottom_border(paragraph, color="B5C9D7", size="8", space="5"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.first_child_found_in("w:pBdr")
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)
    borders.append(border)


def _docx_add_section_heading(doc, title):
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(title)
    _docx_set_run(run, size=12, color=ACCENT_COLOR, bold=True)
    _docx_bottom_border(p, color="B8D0DF", size="5", space="2")
    return p


def _docx_add_body(doc, text, size=9.5, color=TEXT_COLOR, after=2):
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    _docx_set_run(run, size=size, color=color)
    return p


def _docx_add_bullet(doc, text):
    from docx.shared import Cm, Pt

    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.38)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2.3)
    p.paragraph_format.line_spacing = 1.09
    run = p.add_run(text)
    _docx_set_run(run, size=9.45, color=TEXT_COLOR)
    return p


def _docx_configure_styles(doc):
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    normal = doc.styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor(*TEXT_COLOR)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.1


def _docx_add_hyperlink(paragraph, text, url):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run_el = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "19608A")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color_el)
    rpr.append(underline)
    run_el.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run_el.append(text_node)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def export_docx(profile: CandidateProfile, evidence: list[ProfileEvidence]) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError as e:
        raise RuntimeError(f"导出 Word 需要安装 python-docx: pip install python-docx ({e})")

    doc = Document()
    _docx_configure_styles(doc)

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    # 姓名
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(1)
    run = name_p.add_run(_field(profile, "name") or "姓名")
    _docx_set_run(run, size=21, color=TEXT_COLOR, bold=True)

    # 目标岗位
    if _field(profile, "target_role"):
        target_p = doc.add_paragraph()
        target_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        target_p.paragraph_format.space_after = Pt(3)
        run = target_p.add_run(f"求职意向：{_field(profile, 'target_role')}")
        _docx_set_run(run, size=10.5, color=ACCENT_COLOR, bold=True)

    # 联系方式
    contact_parts = []
    if _field(profile, "phone"):
        contact_parts.append(_field(profile, "phone"))
    if _field(profile, "email"):
        contact_parts.append(_field(profile, "email"))
    if _field(profile, "city"):
        contact_parts.append(_field(profile, "city"))
    if contact_parts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(7)
        run = contact_p.add_run("  |  ".join(contact_parts))
        _docx_set_run(run, size=9.2, color=MUTED_COLOR)
        if _field(profile, "homepage"):
            contact_p.add_run("  |  ")
            _docx_add_hyperlink(contact_p, _field(profile, "homepage"), _field(profile, "homepage"))
        _docx_bottom_border(contact_p, color="6E98B2", size="9", space="6")

    # 个人简介
    if _field(profile, "summary"):
        _docx_add_section_heading(doc, "个人概述")
        _docx_add_body(doc, _field(profile, "summary"), size=9.6, after=3)

    # 按 category 分节
    groups = _group_evidence(evidence)
    for cat in CATEGORY_ORDER:
        items = groups.get(cat, [])
        if not items:
            continue
        label = CATEGORY_LABELS.get(cat, cat)
        _docx_add_section_heading(doc, label)
        for item in items:
            if _field(item, "title"):
                title_p = doc.add_paragraph()
                title_p.paragraph_format.space_before = Pt(1)
                title_p.paragraph_format.space_after = Pt(1)
                title_p.paragraph_format.line_spacing = 1.0
                run = title_p.add_run(_field(item, "title"))
                _docx_set_run(run, size=10.2, color=TEXT_COLOR, bold=True)
            for line in _field(item, "content").split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("- ") or line.startswith("· "):
                    _docx_add_bullet(doc, line[2:])
                else:
                    _docx_add_body(doc, line, size=9.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ============================================================
# PDF 导出
# ============================================================

def _pdf_register_fonts():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    system = platform.system()
    if system == "Windows":
        font_dir = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
        candidates = [
            ("msyh.ttc", "MicrosoftYaHei"),
            ("msyhbd.ttc", "MicrosoftYaHeiBold"),
            ("simsun.ttc", "SimSun"),
        ]
    elif system == "Darwin":
        font_dir = Path("/System/Library/Fonts")
        candidates = [
            ("PingFang.ttc", "PingFang"),
            ("STHeiti Light.ttc", "STHeiti"),
        ]
    else:
        font_dir = Path("/usr/share/fonts")
        candidates = [
            ("truetype/wqy/wqy-zenhei.ttc", "WenQuanYi"),
            ("opentype/noto/NotoSansCJK-Regular.ttc", "NotoSansCJK"),
        ]

    registered = set()
    for filename, name in candidates:
        font_path = font_dir / filename
        if font_path.exists() and name not in registered:
            try:
                pdfmetrics.registerFont(TTFont(name, str(font_path)))
                registered.add(name)
            except Exception:
                continue
    return registered


def _pdf_get_fonts():
    registered = _pdf_register_fonts()
    system = platform.system()
    if system == "Windows":
        return "MicrosoftYaHei", "MicrosoftYaHei"
    elif system == "Darwin":
        for name in ["PingFang", "STHeiti"]:
            if name in registered:
                return name, name
    else:
        for name in ["WenQuanYi", "NotoSansCJK"]:
            if name in registered:
                return name, name
    return "Helvetica", "Helvetica"


def export_pdf(profile: CandidateProfile, evidence: list[ProfileEvidence]) -> bytes:
    try:
        from reportlab.lib.colors import HexColor
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, HRFlowable
    except ImportError as e:
        raise RuntimeError(f"导出 PDF 需要安装 reportlab: pip install reportlab ({e})")

    font_cn, _ = _pdf_get_fonts()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=1.25 * cm,
        bottomMargin=1.15 * cm,
        leftMargin=1.7 * cm,
        rightMargin=1.7 * cm,
    )

    styles = getSampleStyleSheet()
    C_TEXT = HexColor("#1C2938")
    C_MUTED = HexColor("#566573")
    C_ACCENT = HexColor("#19597D")

    s_name = ParagraphStyle(
        "ResumeName", parent=styles["Normal"],
        fontName=font_cn, fontSize=20, leading=26,
        alignment=1, spaceAfter=2, textColor=C_TEXT,
    )
    s_target = ParagraphStyle(
        "ResumeTarget", parent=styles["Normal"],
        fontName=font_cn, fontSize=10.5, leading=14,
        alignment=1, spaceAfter=3, textColor=C_ACCENT,
    )
    s_contact = ParagraphStyle(
        "ResumeContact", parent=styles["Normal"],
        fontName=font_cn, fontSize=9, leading=12,
        alignment=1, spaceAfter=6, textColor=C_MUTED,
    )
    s_heading = ParagraphStyle(
        "ResumeHeading", parent=styles["Normal"],
        fontName=font_cn, fontSize=12, leading=16,
        spaceBefore=8, spaceAfter=4, textColor=C_ACCENT,
    )
    s_body = ParagraphStyle(
        "ResumeBody", parent=styles["Normal"],
        fontName=font_cn, fontSize=9.5, leading=13,
        spaceAfter=3, textColor=C_TEXT,
    )
    s_bullet = ParagraphStyle(
        "ResumeBullet", parent=styles["Normal"],
        fontName=font_cn, fontSize=9.5, leading=13,
        leftIndent=12, bulletIndent=0, spaceAfter=2, textColor=C_TEXT,
    )
    s_title = ParagraphStyle(
        "ResumeItemTitle", parent=styles["Normal"],
        fontName=font_cn, fontSize=10, leading=14,
        spaceBefore=2, spaceAfter=1, textColor=C_TEXT,
    )

    story = []

    # 姓名
    story.append(Paragraph(_field(profile, "name") or "姓名", s_name))

    # 目标岗位
    if _field(profile, "target_role"):
        story.append(Paragraph(f"求职意向：{_field(profile, 'target_role')}", s_target))

    # 联系方式
    contact_parts = []
    if _field(profile, "phone"):
        contact_parts.append(_field(profile, "phone"))
    if _field(profile, "email"):
        contact_parts.append(_field(profile, "email"))
    if _field(profile, "city"):
        contact_parts.append(_field(profile, "city"))
    if _field(profile, "homepage"):
        contact_parts.append(
            f'<a href="{_field(profile, "homepage")}" color="#19608A">{_field(profile, "homepage")}</a>'
        )
    if contact_parts:
        story.append(Paragraph("  |  ".join(contact_parts), s_contact))
        story.append(HRFlowable(width="100%", thickness=1, color="#6E98B2", spaceAfter=6))

    # 个人简介
    if _field(profile, "summary"):
        story.append(Paragraph("<b>个人概述</b>", s_heading))
        story.append(HRFlowable(width="100%", thickness=0.5, color="#B8D0DF", spaceAfter=4))
        story.append(Paragraph(_field(profile, "summary"), s_body))

    # 按 category 分节
    groups = _group_evidence(evidence)
    for cat in CATEGORY_ORDER:
        items = groups.get(cat, [])
        if not items:
            continue
        label = CATEGORY_LABELS.get(cat, cat)
        story.append(Paragraph(f"<b>{label}</b>", s_heading))
        story.append(HRFlowable(width="100%", thickness=0.5, color="#B8D0DF", spaceAfter=4))
        for item in items:
            if _field(item, "title"):
                story.append(Paragraph(f"<b>{_field(item, 'title')}</b>", s_title))
            for line in _field(item, "content").split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("- ") or line.startswith("· "):
                    story.append(Paragraph(f"• {line[2:]}", s_bullet))
                else:
                    story.append(Paragraph(line, s_body))

    doc.build(story)
    buf.seek(0)
    return buf.read()
